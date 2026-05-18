#!/usr/bin/env python3
"""Structural lint for proposal documents (.md drafts or .docx finals).

Catches the class of defects that survive Red Team / Gold Team / White Glove
review when they're introduced *after* those reviews run — e.g., a new section
inserted post-review without renumbering (EXTiC 26-2 shipped a whitepaper with
two §4 sections this way).

Checks:
  - Section numbering monotonic — no duplicates, no skips
  - Figure references and captions reconcile (every "Figure N" in prose has a
    "Figure N:" caption, and vice versa)
  - Classification marking present on page 1 (UNCLASSIFIED / CONFIDENTIAL / etc.)
  - Distribution statement present on page 1

Usage:
  # Single .docx (most common — runs against the file you're about to ship)
  python scripts/lint-document-structure.py --docx proposals/extic-26-2/final/docx/whitepaper.docx

  # Single .md (pre-export check)
  python scripts/lint-document-structure.py --md proposals/extic-26-2/drafts/whitepaper.md

  # Every .docx in a proposal's final/docx/ (catches multi-document drift)
  python scripts/lint-document-structure.py --proposal extic-26-2

Exit codes:
  0 = no HIGH findings
  1 = at least one HIGH finding (CI / pre-commit / export-proposal gate)
"""

import argparse
import re
import sys
from collections import Counter
from pathlib import Path

for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
    except Exception:
        pass

WORKSPACE_ROOT = Path(__file__).parent.parent

try:
    from docx import Document  # type: ignore
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(2)


SECTION_RE = re.compile(r"^(\d+(?:\.\d+)*)\.\s+(.+?)$")
FIGURE_CAPTION_RE = re.compile(r"\bFigure\s+(\d+)\s*[:.—–-]\s*(.+?)$", re.IGNORECASE)
FIGURE_REF_RE = re.compile(r"\bFigure\s+(\d+)\b", re.IGNORECASE)
CLASSIFICATION_TOKENS = (
    "UNCLASSIFIED", "CONFIDENTIAL", "SECRET", "TOP SECRET",
    "CUI", "FOUO", "FOR OFFICIAL USE ONLY",
)
DISTRIBUTION_TOKENS = ("Distribution Statement", "DISTRIBUTION STATEMENT")


def extract_headings_and_text_from_docx(path: Path) -> tuple[list[tuple[int, str]], str, str]:
    """Read a .docx and return (headings, full_text, first_page_text)."""
    doc = Document(path)
    headings: list[tuple[int, str]] = []
    all_text: list[str] = []
    first_page: list[str] = []
    # python-docx doesn't expose page breaks directly; use first ~30 paragraphs
    # as a proxy for "first page" content.
    para_count = 0
    for para in doc.paragraphs:
        text = para.text
        all_text.append(text)
        if para_count < 30:
            first_page.append(text)
        para_count += 1
        style = para.style.name if para.style else ""
        if style.startswith("Heading"):
            level_str = style.replace("Heading", "").strip()
            level = int(level_str) if level_str.isdigit() else 1
            headings.append((level, text.strip()))
    # Tables — flatten cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)
    return headings, "\n".join(all_text), "\n".join(first_page)


def extract_headings_and_text_from_md(path: Path) -> tuple[list[tuple[int, str]], str, str]:
    """Read a .md and return (headings, full_text, first_page_text)."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    raw = re.sub(r"<!--.*?-->", "", raw, flags=re.DOTALL)  # strip HTML comments
    headings: list[tuple[int, str]] = []
    for line in raw.split("\n"):
        m = re.match(r"^(#+)\s+(.+?)$", line)
        if m:
            headings.append((len(m.group(1)), m.group(2).strip()))
    # First page proxy = first ~30 lines that aren't blank
    first_page_lines = [ln for ln in raw.split("\n") if ln.strip()][:30]
    return headings, raw, "\n".join(first_page_lines)


def lint_section_numbering(headings: list[tuple[int, str]]) -> list[dict]:
    """Detect duplicate or skipped section numbers."""
    findings: list[dict] = []
    # Pull numbered headings only (level 1-2 with leading "N." or "N.M.")
    numbered: list[tuple[str, str]] = []  # (number_str, text)
    for level, text in headings:
        if level > 3:
            continue
        m = SECTION_RE.match(text)
        if m:
            numbered.append((m.group(1), text))

    # Top-level (whole-number) sequence
    top_level = [n for n, _ in numbered if "." not in n]
    counts = Counter(top_level)
    for num, c in counts.items():
        if c > 1:
            duplicates = [t for n, t in numbered if n == num]
            findings.append({
                "severity": "HIGH",
                "category": "section-numbering",
                "issue": f"Section number {num} appears {c} times",
                "detail": " | ".join(duplicates),
            })

    # Monotonic check on top-level
    seen_ints = []
    for n in top_level:
        try:
            seen_ints.append(int(n))
        except ValueError:
            pass
    if seen_ints:
        for i in range(1, len(seen_ints)):
            if seen_ints[i] < seen_ints[i - 1]:
                findings.append({
                    "severity": "MEDIUM",
                    "category": "section-numbering",
                    "issue": f"Section {seen_ints[i]} appears after section {seen_ints[i - 1]} (out of order)",
                    "detail": "",
                })
            elif seen_ints[i] > seen_ints[i - 1] + 1:
                # Skip — only flag once per gap, only if we expected sequential
                gap = seen_ints[i] - seen_ints[i - 1] - 1
                findings.append({
                    "severity": "MEDIUM",
                    "category": "section-numbering",
                    "issue": f"Section numbering skips {gap} integer(s) between {seen_ints[i - 1]} and {seen_ints[i]}",
                    "detail": "",
                })

    return findings


def lint_figures(full_text: str) -> list[dict]:
    """Reconcile figure references against captions."""
    findings: list[dict] = []
    captions: set[int] = set()
    refs: set[int] = set()
    for line in full_text.split("\n"):
        # A caption is a line that BOTH names "Figure N" AND has descriptive text after it.
        m = FIGURE_CAPTION_RE.search(line)
        if m:
            captions.add(int(m.group(1)))
        # Refs are inline mentions of "Figure N" in narrative
        for rm in FIGURE_REF_RE.finditer(line):
            refs.add(int(rm.group(1)))

    if not captions and not refs:
        return findings

    orphan_refs = refs - captions
    orphan_caps = captions - refs

    for n in sorted(orphan_refs):
        findings.append({
            "severity": "HIGH",
            "category": "figures",
            "issue": f"Figure {n} referenced in prose but no caption found",
            "detail": "",
        })
    for n in sorted(orphan_caps):
        # Captions without refs are softer — sometimes graphics speak for themselves.
        findings.append({
            "severity": "LOW",
            "category": "figures",
            "issue": f"Figure {n} has a caption but no prose reference",
            "detail": "",
        })
    return findings


def lint_first_page_markings(first_page: str) -> list[dict]:
    """Classification + distribution statement on page 1."""
    findings: list[dict] = []
    upper = first_page.upper()
    if not any(tok in upper for tok in CLASSIFICATION_TOKENS):
        findings.append({
            "severity": "HIGH",
            "category": "classification",
            "issue": "No classification marking detected on page 1 (looked for UNCLASSIFIED / CONFIDENTIAL / SECRET / CUI / FOUO)",
            "detail": "",
        })
    if not any(tok in first_page for tok in DISTRIBUTION_TOKENS):
        findings.append({
            "severity": "MEDIUM",
            "category": "distribution",
            "issue": "No distribution statement detected on page 1",
            "detail": "",
        })
    return findings


def lint_one(path: Path) -> list[dict]:
    if path.suffix.lower() == ".docx":
        headings, full, first_page = extract_headings_and_text_from_docx(path)
    elif path.suffix.lower() == ".md":
        headings, full, first_page = extract_headings_and_text_from_md(path)
    else:
        return [{"severity": "HIGH", "category": "input", "issue": f"Unsupported extension: {path.suffix}", "detail": ""}]

    findings = []
    findings.extend(lint_section_numbering(headings))
    findings.extend(lint_figures(full))
    findings.extend(lint_first_page_markings(first_page))
    return findings


def print_report(path: Path, findings: list[dict]) -> int:
    """Return number of HIGH findings."""
    print(f"\n── {path} ──")
    if not findings:
        print("  ✓ no issues found")
        return 0
    high = 0
    for f in findings:
        sev = f["severity"]
        marker = "✗" if sev == "HIGH" else ("⚠" if sev == "MEDIUM" else "·")
        line = f"  {marker} [{sev:6}] [{f['category']}] {f['issue']}"
        if f["detail"]:
            line += f"\n      {f['detail']}"
        print(line)
        if sev == "HIGH":
            high += 1
    return high


def collect_targets(args: argparse.Namespace) -> list[Path]:
    targets: list[Path] = []
    if args.docx:
        targets.append(Path(args.docx))
    if args.md:
        targets.append(Path(args.md))
    if args.proposal:
        prop_dir = WORKSPACE_ROOT / "proposals" / args.proposal
        if not prop_dir.exists():
            print(f"ERROR: proposal not found at {prop_dir}", file=sys.stderr)
            sys.exit(2)
        targets.extend(sorted(
            f for f in (prop_dir / "final" / "docx").glob("*.docx")
            if not f.name.startswith("~$")  # skip Word owner-lock stubs (not valid .docx)
        ))
        if not targets:
            # Fall back to drafts if no .docx exists yet
            targets.extend(sorted((prop_dir / "drafts").glob("*.md")))
    return targets


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--docx", help="Single .docx to lint")
    ap.add_argument("--md", help="Single .md to lint")
    ap.add_argument("--proposal", help="Lint every final/docx/*.docx in this proposal (fall back to drafts/ if no finals)")
    args = ap.parse_args()

    if not (args.docx or args.md or args.proposal):
        ap.print_help()
        return 1

    targets = collect_targets(args)
    if not targets:
        print("No documents found to lint.", file=sys.stderr)
        return 1

    total_high = 0
    for path in targets:
        if not path.exists():
            print(f"ERROR: {path} not found", file=sys.stderr)
            total_high += 1
            continue
        findings = lint_one(path)
        total_high += print_report(path, findings)

    print()
    print(f"Done — {len(targets)} document(s) checked, {total_high} HIGH finding(s)")
    return 0 if total_high == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
