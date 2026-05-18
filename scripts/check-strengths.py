#!/usr/bin/env python3
"""Verify Gold-Team Significant Strengths survive into the final document.

Parses each Significant Strength from the most recent Gold Team scorecard,
extracts the distinctive phrases that make it a Strength (proper nouns,
alphanumeric codes, quantified claims, specific named entities), then
searches the current draft (or final .docx) for those phrases.

For each Strength, reports:
  PRESENT  — most distinctive phrases present in the draft (>=70%)
  REDUCED  — partial coverage (30–70%)
  MISSING  — most phrases absent (<30%)

This is the check that would have caught the EXTiC 26-2 MOS-adapter-slate
stripping: Gold Team called out "FA14 anchor at D+30, 14E/14G/14H/14P + 17E
at D+60/D+90/D+120, Fort Sill / Fort Eisenhower, FM 3-01 / ATP 3-01.7/8 /
FM 3-12 / JP 3-13.1" as the decisive Significant Strength. The submitted
whitepaper kept only "FA14 Air Defense model configuration" — every other
distinctive phrase was stripped during team review.

Usage:
  # Check drafts (pre-export)
  python scripts/check-strengths.py --proposal extic-26-2

  # Check a specific .docx (post-team-revision, before submit)
  python scripts/check-strengths.py --proposal extic-26-2 --target docx

  # Custom target path
  python scripts/check-strengths.py --proposal extic-26-2 --target-path path/to/file.docx

Output:
  - Console report ranked by status
  - reviews/strength-preservation.md (regenerated each run)

Exit codes:
  0 = no MISSING strengths (all PRESENT or REDUCED is acceptable; Bill triages)
  1 = at least one MISSING strength (CI / pre-submit gate)
"""

import argparse
import re
import sys
from datetime import date
from pathlib import Path

for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[attr-defined]
    except Exception:
        pass

WORKSPACE_ROOT = Path(__file__).parent.parent

try:
    from docx import Document  # type: ignore
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(2)


# Stopwords — generic words/phrases that shouldn't be counted as load-bearing
# distinctive phrases even though they're capitalized or look specific.
STOPWORDS = {
    "Phase 1", "Phase 2", "Phase 3", "Phase 4",
    "Government", "Federal", "DoD", "DoW", "Department",
    "Significant Strength", "Significant Weakness", "Strength", "Weakness",
    "Basis", "Benefit", "Risk", "Fix", "Note",
    "Section", "Sections", "Paragraph", "Table", "Figure",
    "AI", "API", "OS",
    "TRL", "MOS", "OEM", "TTP", "TTPs", "USA", "USAF",
    "Acme", "Acme Defense",  # the company name appears everywhere; not a discriminator
}


def extract_distinctive_phrases(text: str) -> set[str]:
    """Pull capitalized multi-word phrases, alphanumeric codes, and quantified claims."""
    text = re.sub(r"<!--.*?-->", "", text)
    text = re.sub(r"`[^`]+`", " ", text)  # strip code spans
    phrases: set[str] = set()

    # 1. Alphanumeric codes with internal digits/dashes: FA14, IL-6, GPT-5, 14E, 17E, JP 3-13.1, AJP-01
    for m in re.finditer(r"\b[A-Z][A-Z0-9]*[-/]?[A-Z0-9.]*\d[A-Z0-9./-]*\b", text):
        s = m.group(0)
        if len(s) >= 2 and not s.isdigit():
            phrases.add(s)
    # Also catch patterns like "FM 3-01", "ATP 3-01.7/8" (letter-prefix space number)
    for m in re.finditer(r"\b[A-Z]{2,4}\s+\d+[\d./-]*\b", text):
        phrases.add(m.group(0))

    # 2. Multi-word capitalized phrases (proper-noun-like): "Fort Sill", "Air Defense Artillery", "Navy ICOP"
    for m in re.finditer(r"\b[A-Z][a-z]+(?:\s+(?:[A-Z][a-zA-Z0-9.-]+|of|and|the|&)){1,4}\b", text):
        s = m.group(0).strip()
        # Strip trailing connector words
        s = re.sub(r"\s+(?:of|and|the|&)$", "", s)
        # Require at least one fully-capitalized word
        words = s.split()
        if len(words) >= 2 and any(w[0].isupper() and w not in {"of", "and", "the"} for w in words):
            phrases.add(s)

    # 3. Standalone ALL-CAPS acronyms (3+ chars): USSOCOM, JIATF, TRADOC, CENTCOM, NIWC
    for m in re.finditer(r"\b[A-Z]{3,}\b", text):
        phrases.add(m.group(0))

    # 4. Quantified claims: percentages, "D+30", "TRL 8", "IL-6"
    for m in re.finditer(r"\b\d+(?:\.\d+)?%", text):
        phrases.add(m.group(0))
    for m in re.finditer(r"\bD\+\d+\b", text):
        phrases.add(m.group(0))
    for m in re.finditer(r"\bTRL\s?\d\b", text):
        phrases.add(m.group(0))
    for m in re.finditer(r"\bIL[-\s]?\d\b", text):
        phrases.add(m.group(0))

    # Filter stopwords
    return {p for p in phrases if p not in STOPWORDS and len(p) > 1}


def _clean_md(s: str) -> str:
    """Strip markdown emphasis (**, *), leading list markers, surrounding whitespace."""
    s = re.sub(r"\*+", "", s)
    s = re.sub(r"^[-*\s]+", "", s)
    return s.strip().rstrip(".")


def extract_significant_strengths(scorecard_text: str) -> list[dict]:
    """Pull each Significant Strength block (header + basis + benefit) from a Gold Team scorecard."""
    if not scorecard_text:
        return []
    strengths = []
    pattern = re.compile(
        r"^#{1,6}\s+(Significant Strength[^\n]*)\n+"
        r"(?:[-*]\s*)?\*{0,2}Basis\*{0,2}:\s*(.+?)\n"
        r"(?:[-*]\s*)?\*{0,2}Benefit[^:]*\*{0,2}:\s*(.+?)\n",
        re.MULTILINE | re.DOTALL,
    )
    for m in pattern.finditer(scorecard_text):
        title = _clean_md(m.group(1))
        basis = _clean_md(m.group(2).split("\n")[0])
        benefit = _clean_md(m.group(3).split("\n")[0])
        strengths.append({"title": title, "basis": basis, "benefit": benefit})
    return strengths


def read_target_text(prop_dir: Path, target_mode: str, target_path: str | None) -> tuple[str, str]:
    """Return (text, source_description)."""
    if target_path:
        p = Path(target_path)
        if not p.exists():
            print(f"ERROR: {p} not found", file=sys.stderr)
            sys.exit(2)
        return _read_any(p), str(p)
    if target_mode == "drafts":
        parts = []
        for md in sorted((prop_dir / "drafts").glob("*.md")):
            parts.append(md.read_text(encoding="utf-8", errors="replace"))
        return "\n\n".join(parts), str(prop_dir / "drafts")
    if target_mode == "docx":
        parts = []
        docx_dir = prop_dir / "final" / "docx"
        for d in sorted(f for f in docx_dir.glob("*.docx") if not f.name.startswith("~$")):
            parts.append(_read_any(d))
        if not parts:
            print(f"ERROR: no .docx files in {docx_dir}", file=sys.stderr)
            sys.exit(2)
        return "\n\n".join(parts), str(docx_dir)
    raise ValueError(f"Unknown target mode: {target_mode}")


def _read_any(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def classify(strength: dict, target_text: str) -> dict:
    """Determine PRESENT / REDUCED / MISSING status for one strength."""
    # Build the "must-have" phrase set from the strength's basis + benefit text
    must_have = extract_distinctive_phrases(strength["basis"] + " " + strength["benefit"])
    if not must_have:
        return {**strength, "status": "PRESENT", "score": 1.0, "found": set(), "missing": set(), "note": "no distinctive phrases extracted from strength text"}

    found: set[str] = set()
    missing: set[str] = set()
    target_lower = target_text.lower()
    for phrase in must_have:
        # Case-insensitive substring match for robustness against minor formatting drift
        if phrase.lower() in target_lower:
            found.add(phrase)
        else:
            missing.add(phrase)

    score = len(found) / len(must_have) if must_have else 1.0
    if score >= 0.7:
        status = "PRESENT"
    elif score >= 0.3:
        status = "REDUCED"
    else:
        status = "MISSING"

    return {**strength, "status": status, "score": score, "found": found, "missing": missing, "note": ""}


def render_report(slug: str, source_desc: str, results: list[dict]) -> str:
    counts = {"PRESENT": 0, "REDUCED": 0, "MISSING": 0}
    for r in results:
        counts[r["status"]] += 1

    out = []
    out.append(f"# Gold-Team Strength Preservation Report — {slug}")
    out.append("")
    out.append(f"**Date:** {date.today().isoformat()}")
    out.append(f"**Target:** `{source_desc}`")
    out.append(f"**Strengths checked:** {len(results)}")
    out.append("")
    out.append(f"- ✓ **PRESENT:** {counts['PRESENT']}")
    out.append(f"- ⚠ **REDUCED:** {counts['REDUCED']}")
    out.append(f"- ✗ **MISSING:** {counts['MISSING']}")
    out.append("")
    out.append("Verdict: " + (
        "**MISSING strengths — pre-submit gate FAILED.** Restore or document the omissions before submission."
        if counts["MISSING"]
        else ("**REDUCED strengths — review.** Some distinctive phrases dropped during revision. Confirm intentional."
              if counts["REDUCED"]
              else "**All strengths preserved.** Submission-safe from this check's perspective.")
    ))
    out.append("")
    out.append("---")
    out.append("")

    for status_order in ("MISSING", "REDUCED", "PRESENT"):
        for r in results:
            if r["status"] != status_order:
                continue
            marker = {"PRESENT": "✓", "REDUCED": "⚠", "MISSING": "✗"}[r["status"]]
            out.append(f"## {marker} [{r['status']}] {r['title']}")
            out.append("")
            out.append(f"- **Cited basis:** {r['basis']}")
            out.append(f"- **Why it matters (Gold Team):** {r['benefit']}")
            out.append(f"- **Coverage score:** {int(r['score'] * 100)}%")
            if r["found"]:
                shown = sorted(r["found"])[:20]
                out.append(f"- **Phrases preserved ({len(r['found'])}):** " + ", ".join(f"`{p}`" for p in shown) +
                           (" …" if len(r["found"]) > 20 else ""))
            if r["missing"]:
                shown = sorted(r["missing"])[:20]
                out.append(f"- **Phrases MISSING ({len(r['missing'])}):** " + ", ".join(f"`{p}`" for p in shown) +
                           (" …" if len(r["missing"]) > 20 else ""))
            if r["note"]:
                out.append(f"- **Note:** {r['note']}")
            if r["status"] == "MISSING":
                out.append("- **Recommended action:** restore the named entities listed under MISSING, or document why the strength can no longer be claimed.")
            elif r["status"] == "REDUCED":
                out.append("- **Recommended action:** review the MISSING phrases; if any are load-bearing (named program, quantified claim, named partner), restore them.")
            out.append("")

    return "\n".join(out).rstrip() + "\n"


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--proposal", required=True, help="Proposal slug")
    ap.add_argument("--target", choices=("drafts", "docx"), default="drafts",
                    help="What to check against (default: drafts/)")
    ap.add_argument("--target-path", help="Override: specific .md or .docx file to check")
    ap.add_argument("--out", help="Output report path (default: proposals/<slug>/reviews/strength-preservation.md)")
    args = ap.parse_args()

    prop_dir = WORKSPACE_ROOT / "proposals" / args.proposal
    if not prop_dir.exists():
        print(f"ERROR: proposal not found at {prop_dir}", file=sys.stderr)
        return 2

    scorecard_path = prop_dir / "reviews" / "gold-team-scorecard.md"
    if not scorecard_path.exists():
        print(f"ERROR: gold-team-scorecard.md not found at {scorecard_path}", file=sys.stderr)
        print("Run /red-team-review --mode=gold first.", file=sys.stderr)
        return 2

    strengths = extract_significant_strengths(scorecard_path.read_text(encoding="utf-8", errors="replace"))
    if not strengths:
        print(f"WARN: no Significant Strengths found in {scorecard_path}.", file=sys.stderr)
        print("Either Gold Team is still in progress or the parser couldn't match the headers.", file=sys.stderr)
        return 0

    target_text, source_desc = read_target_text(prop_dir, args.target, args.target_path)

    results = [classify(s, target_text) for s in strengths]
    report = render_report(args.proposal, source_desc, results)

    out_path = Path(args.out) if args.out else prop_dir / "reviews" / "strength-preservation.md"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(report, encoding="utf-8")
    print(f"Wrote {out_path}")

    # Console summary
    print()
    counts = {"PRESENT": 0, "REDUCED": 0, "MISSING": 0}
    for r in results:
        counts[r["status"]] += 1
        marker = {"PRESENT": "✓", "REDUCED": "⚠", "MISSING": "✗"}[r["status"]]
        print(f"  {marker} [{r['status']:7}] {r['title']}  ({int(r['score'] * 100)}%)")
    print()
    print(f"Summary: ✓ PRESENT {counts['PRESENT']}  ⚠ REDUCED {counts['REDUCED']}  ✗ MISSING {counts['MISSING']}")

    return 1 if counts["MISSING"] else 0


if __name__ == "__main__":
    sys.exit(main())
