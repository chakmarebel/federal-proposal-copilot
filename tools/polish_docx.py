#!/usr/bin/env python3
"""
polish_docx.py - post-export polish for proposal Word documents.

Runs after tools/md_to_docx.py. For every .docx in a proposal's final/docx/:
  1. Optimizes table column widths - narrow short-content columns, widen wordy
     ones - instead of the converter's even split.
  2. Adds a running header (the document's first H1 title, pages 2+).
  3. Adds a footer with the company name and a "Page X of Y" field (all pages).

Usage:
  python tools/polish_docx.py --proposal <slug> [--company "[Your Company], Inc."]
  python tools/polish_docx.py --files path/to/a.docx path/to/b.docx

Idempotent: re-running re-applies the same polish. If you re-export from
markdown, re-run this script - the export step does not preserve it.
"""
import argparse, glob, os, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

USABLE_WIDTH = 6.5   # inches: 8.5" Letter minus 1" margins each side
MIN_COL = 0.5        # inches: minimum column width
WEIGHT_CAP = 40      # cap per-column weight so one very long column cannot
                     # starve the medium-content columns next to it


def optimize_table(table):
    ncols = len(table.columns)
    if ncols == 0:
        return
    maxlen = [1] * ncols
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < ncols:
                maxlen[i] = max(maxlen[i], len(cell.text.strip()))
    weights = [min(m, WEIGHT_CAP) for m in maxlen]
    total = float(sum(weights))
    widths = [USABLE_WIDTH * w / total for w in weights]
    # clamp to minimum, redistribute the deficit across the roomier columns
    deficit = 0.0
    free = []
    for i, w in enumerate(widths):
        if w < MIN_COL:
            deficit += MIN_COL - w
            widths[i] = MIN_COL
        else:
            free.append(i)
    if deficit > 0 and free:
        freetotal = sum(widths[i] for i in free)
        for i in free:
            widths[i] -= deficit * widths[i] / freetotal

    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    # set the table grid columns
    grid = tbl.find(qn('w:tblGrid'))
    if grid is not None:
        for i, gridCol in enumerate(grid.findall(qn('w:gridCol'))):
            if i < ncols:
                gridCol.set(qn('w:w'), str(int(widths[i] * 1440)))
    # set every cell width so Word honors the fixed layout
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < ncols:
                cell.width = Inches(widths[i])
    return widths


def _page_field(paragraph, instr):
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
    t = OxmlElement('w:instrText'); t.set(qn('xml:space'), 'preserve'); t.text = instr
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
    run._r.append(b); run._r.append(t); run._r.append(e)
    run.font.size = Pt(9)
    return run


def first_title(doc):
    for p in doc.paragraphs:
        sn = (p.style.name or '') if p.style else ''
        if sn.startswith('Heading 1') or sn == 'Title':
            if p.text.strip():
                return p.text.strip()
    for p in doc.paragraphs:
        if p.text.strip():
            return p.text.strip()
    return ''


def style_runs(paragraph, size=9, grey=True):
    for r in paragraph.runs:
        r.font.size = Pt(size)
        if grey:
            r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def add_header_footer(doc, company):
    title = first_title(doc)
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # header on pages 2+ only (page 1 already shows the H1 title)
        hp = section.header.paragraphs[0]
        hp.text = title
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_runs(hp)
        section.first_page_header.paragraphs[0].text = ''

        # footer on every page: company name, then Page X of Y
        for fp in (section.footer.paragraphs[0],
                   section.first_page_footer.paragraphs[0]):
            fp.text = ''
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run(company + '      |      Page ')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            _page_field(fp, 'PAGE')
            of = fp.add_run(' of ')
            of.font.size = Pt(9)
            of.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            _page_field(fp, 'NUMPAGES')


def polish(path, company):
    doc = Document(path)
    ntables = 0
    for table in doc.tables:
        optimize_table(table)
        ntables += 1
    add_header_footer(doc, company)
    doc.save(path)
    print(f'  [OK] {os.path.basename(path)} - {ntables} table(s) optimized, header/footer added')


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--proposal')
    ap.add_argument('--files', nargs='*')
    ap.add_argument('--company', default='[Your Company], Inc.')
    args = ap.parse_args()

    if args.files:
        targets = args.files
    elif args.proposal:
        root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        targets = sorted(glob.glob(os.path.join(
            root, 'proposals', args.proposal, 'final', 'docx', '*.docx')))
    else:
        ap.error('pass --proposal <slug> or --files <paths>')

    if not targets:
        print('No .docx files found.'); sys.exit(1)
    print(f'Polishing {len(targets)} document(s):')
    for t in targets:
        polish(t, args.company)
    print('Done.')


if __name__ == '__main__':
    main()
