#!/usr/bin/env python3
"""
tools/md_to_docx.py — Shared markdown-to-Word converter for federal-proposal-assistant.

Converts proposal draft .md files → .docx in proposals/<slug>/final/docx/.

Usage:
    # All drafts, both individual + combined (default):
    python tools/md_to_docx.py --proposal nato-diana-decision-superiority

    # Individual files only:
    python tools/md_to_docx.py --proposal myproposal --mode individual

    # Combined document only, custom name:
    python tools/md_to_docx.py --proposal myproposal --mode combined --combined-name "technical-volume"

    # Explicit ordered file list (filenames relative to drafts/):
    python tools/md_to_docx.py --proposal myproposal --files sec1.md sec2.md sec3.md

    # Explicit workspace root (if running from a different working directory):
    python tools/md_to_docx.py --proposal myproposal --workspace "C:/path/to/federal-proposal-assistant"

Supported markdown:
    # H1  ## H2  ### H3  #### H4  ##### H5  ###### H6
    **bold**  *italic*  ***bold-italic***  `code` (inline monospace)
    - bullets (with nested indent via leading spaces)
    1. numbered lists (with nested indent)
    | table | rows | (pipe-delimited, first row = bold header)
    > blockquote (indented, used for action captions)
    ``` fenced code blocks ``` (verbatim, monospace, auto-sized to fit column)
    <!-- figure: NAME --> (embeds graphics/rendered/NAME.png, centered, fit to column)
    --- horizontal rule (skipped — internal metadata separator)
    <!-- html comments --> (stripped, except figure markers above)
    *Note: ... lines (skipped — internal metadata)
"""

import argparse
import os
import re
import sys
from pathlib import Path

# ── Dependency check ─────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── Constants ─────────────────────────────────────────────────────────────────
WORKSPACE_ROOT = Path(__file__).parent.parent  # tools/ → workspace root
BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
MARGIN_TOP = Inches(1.0)
MARGIN_BOTTOM = Inches(1.0)
MARGIN_LEFT = Inches(1.25)
MARGIN_RIGHT = Inches(1.25)
TEXT_WIDTH_INCHES = 8.5 - 1.25 - 1.25  # 6.0 inches usable width


# ── Document setup ────────────────────────────────────────────────────────────

def setup_document(doc: Document) -> None:
    """Apply margins and default body font to every section."""
    for section in doc.sections:
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

    # Set default run font in Normal style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE


# ── Inline formatting ─────────────────────────────────────────────────────────

def add_formatted_runs(para, text: str) -> None:
    """
    Add text to a paragraph with inline formatting:
        ***bold-italic***  **bold**  *italic*  `code`
    """
    # Split on any inline marker
    pattern = re.compile(
        r'(\*\*\*[^*]+\*\*\*'   # ***bold-italic***
        r'|\*\*[^*]+\*\*'       # **bold**
        r'|\*[^*]+\*'           # *italic*
        r'|`[^`]+`)'            # `code`
    )
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        run = para.add_run()
        run.font.name = BODY_FONT
        if re.match(r'^\*\*\*[^*]+\*\*\*$', part):
            run.text = part[3:-3]
            run.bold = True
            run.italic = True
        elif re.match(r'^\*\*[^*]+\*\*$', part):
            run.text = part[2:-2]
            run.bold = True
        elif re.match(r'^\*[^*]+\*$', part):
            run.text = part[1:-1]
            run.italic = True
        elif re.match(r'^`[^`]+`$', part):
            run.text = part[1:-1]
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            run.text = part


# ── Table parsing ─────────────────────────────────────────────────────────────

def _is_separator_row(line: str) -> bool:
    """True for lines like |---|:---|:---:| that separate header from body."""
    return bool(re.match(r'^\s*\|[\s\-:|]+\|\s*$', line.rstrip()))


def _is_table_line(line: str) -> bool:
    return line.strip().startswith('|')


def _parse_cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip('|').split('|')]


def process_table(doc: Document, table_lines: list[str]) -> None:
    """Render a markdown pipe-table as a Word table."""
    data_rows = [ln for ln in table_lines if not _is_separator_row(ln)]
    if not data_rows:
        return

    num_cols = len(_parse_cells(data_rows[0]))
    if num_cols == 0:
        return

    col_width = Inches(TEXT_WIDTH_INCHES / num_cols)

    table = doc.add_table(rows=len(data_rows), cols=num_cols)
    table.style = "Table Grid"

    for r_idx, row_line in enumerate(data_rows):
        cells = _parse_cells(row_line)
        is_header = r_idx == 0
        for c_idx in range(num_cols):
            cell_text = cells[c_idx] if c_idx < len(cells) else ""
            cell = table.rows[r_idx].cells[c_idx]
            # Set column width
            cell.width = col_width
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            add_formatted_runs(para, cell_text)
            if is_header:
                for run in para.runs:
                    run.bold = True

    # Add a blank paragraph after the table for breathing room
    doc.add_paragraph()


# ── List helpers ──────────────────────────────────────────────────────────────

def _detect_list_indent(line: str) -> int:
    """Return nesting level 0,1,2,... based on leading spaces (every 2 spaces = 1 level)."""
    stripped = line.lstrip(' \t')
    leading = len(line) - len(stripped)
    return min(leading // 2, 8)  # cap at 8 levels


def _add_list_para(doc: Document, text: str, style: str, level: int) -> None:
    """Add a bullet or numbered list paragraph at the given nesting level."""
    para = doc.add_paragraph(style=style)
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    add_formatted_runs(para, text)


# ── Code block ──────────────────────────────────────────────────────────────

def process_code_block(doc: Document, code_lines: list[str]) -> None:
    """Render a fenced code block verbatim in a monospace font.

    ASCII diagrams and pre-formatted tables only align in a fixed-width font.
    Font size is scaled down so the widest line fits the 6.0" text column.
    """
    if not code_lines:
        return
    max_len = max((len(ln) for ln in code_lines), default=1) or 1
    # Courier New advance width is ~0.6 em. Solve max_len * pt * 0.6 / 72 <= width.
    fit_pt = int((TEXT_WIDTH_INCHES * 72) / (0.6 * max_len))
    size = Pt(max(5, min(9, fit_pt)))
    for ln in code_lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(ln)  # verbatim — no inline-markdown parsing
        run.font.name = "Courier New"
        run.font.size = size
    doc.add_paragraph()  # breathing room after the block


# ── Figure embedding ────────────────────────────────────────────────────────

def add_figure(doc: Document, md_path: Path, fig_name: str) -> None:
    """Embed a rendered proposal figure referenced by `<!-- figure: NAME -->`.

    NAME resolves to `<proposal>/graphics/rendered/NAME.png` (the path
    /proposal-graphics renders to). The figure is centered and scaled to the
    text-column width, with height capped to one page so a tall graphic does
    not overflow. A missing PNG renders a visible placeholder rather than
    silently dropping the figure.
    """
    png = (md_path.parent.parent / "graphics" / "rendered" / f"{fig_name}.png").resolve()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if not png.exists():
        run = para.add_run(f"[figure not found: {fig_name}]")
        run.italic = True
        run.font.name = BODY_FONT
        return
    w_in = TEXT_WIDTH_INCHES
    try:
        from PIL import Image as _PILImage
        with _PILImage.open(png) as im:
            iw, ih = im.size
        if iw and ih and (TEXT_WIDTH_INCHES * ih / iw) > 8.0:
            w_in = 8.0 * iw / ih  # height-constrained — keep within one page
    except Exception:
        pass  # PIL absent — width-only scaling still preserves aspect ratio
    para.add_run().add_picture(str(png), width=Inches(w_in))
    doc.add_paragraph()  # breathing room after the figure


# ── Core converter ────────────────────────────────────────────────────────────

def convert_md_to_doc(md_path: Path, doc: Document, page_break: bool = False) -> int:
    """
    Parse one markdown file and append its content to doc.
    Returns the number of non-blank lines processed.
    """
    raw = md_path.read_text(encoding="utf-8")
    # Strip HTML comments — but keep `<!-- figure: ... -->` markers, which the
    # loop below turns into embedded images.
    raw = re.sub(r'<!--(?!\s*figure:).*?-->', '', raw, flags=re.DOTALL)
    lines = raw.split('\n')

    if page_break:
        doc.add_page_break()

    line_count = 0
    i = 0
    while i < len(lines):
        raw_line = lines[i]
        s = raw_line.rstrip()
        stripped = s.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Skip internal metadata
        if stripped.startswith('*Note') or re.match(r'^-{3,}$', stripped):
            i += 1
            continue

        # Fenced code block ``` ... ``` — render verbatim in a monospace font
        if stripped.startswith('```'):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].lstrip().startswith('```'):
                code_lines.append(lines[i])
                line_count += 1
                i += 1
            i += 1  # consume the closing fence (or step past EOF if unclosed)
            process_code_block(doc, code_lines)
            continue

        # H1
        if re.match(r'^# [^#]', stripped):
            doc.add_heading(stripped[2:].strip(), level=1)
            line_count += 1; i += 1; continue

        # H2
        if re.match(r'^## [^#]', stripped):
            doc.add_heading(stripped[3:].strip(), level=2)
            line_count += 1; i += 1; continue

        # H3
        if re.match(r'^### [^#]', stripped):
            doc.add_heading(stripped[4:].strip(), level=3)
            line_count += 1; i += 1; continue

        # H4
        if re.match(r'^#### [^#]', stripped):
            doc.add_heading(stripped[5:].strip(), level=4)
            line_count += 1; i += 1; continue

        # H5
        if re.match(r'^##### [^#]', stripped):
            doc.add_heading(stripped[6:].strip(), level=5)
            line_count += 1; i += 1; continue

        # H6
        if re.match(r'^###### [^#]', stripped):
            doc.add_heading(stripped[7:].strip(), level=6)
            line_count += 1; i += 1; continue

        # Table — collect consecutive table lines
        if _is_table_line(stripped):
            tbl = []
            while i < len(lines) and _is_table_line(lines[i].rstrip()):
                tbl.append(lines[i].rstrip())
                i += 1
            process_table(doc, tbl)
            line_count += len(tbl)
            continue

        # Figure marker — <!-- figure: NAME --> on its own line
        fig_m = re.match(r'^<!--\s*figure:\s*([A-Za-z0-9._-]+)\s*-->$', stripped)
        if fig_m:
            add_figure(doc, md_path, fig_m.group(1))
            line_count += 1; i += 1; continue

        # Bullet list (- or * prefix, possibly indented)
        bullet_m = re.match(r'^(\s*)[-*] (.+)', s)
        if bullet_m:
            level = _detect_list_indent(s)
            _add_list_para(doc, bullet_m.group(2).strip(), "List Bullet", level)
            line_count += 1; i += 1; continue

        # Numbered list (possibly indented)
        num_m = re.match(r'^(\s*)(\d+)\. (.+)', s)
        if num_m:
            level = _detect_list_indent(s)
            _add_list_para(doc, num_m.group(3).strip(), "List Number", level)
            line_count += 1; i += 1; continue

        # Blockquote (> ...) — used for action captions
        if stripped.startswith('> '):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Style it italic
            run = para.add_run(stripped[2:].strip())
            run.italic = True
            run.font.name = BODY_FONT
            run.font.size = Pt(10)
            line_count += 1; i += 1; continue

        # Regular paragraph with inline formatting
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        add_formatted_runs(para, stripped)
        line_count += 1; i += 1

    return line_count


# ── Discovery ─────────────────────────────────────────────────────────────────

def discover_drafts(drafts_dir: Path) -> list[Path]:
    """
    Return all .md files in drafts_dir sorted alphabetically.
    This matches the naming convention (01-..., 02-..., etc.) used in proposal drafts.
    """
    files = sorted(drafts_dir.glob("*.md"))
    return files


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convert proposal markdown drafts to Word .docx files.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("--proposal", required=True,
                        help="Proposal slug (directory name under proposals/)")
    parser.add_argument("--mode", choices=["individual", "combined", "both"],
                        default="both",
                        help="Output mode (default: both)")
    parser.add_argument("--combined-name", default="full-proposal-combined",
                        help="Stem for the combined .docx file (default: full-proposal-combined)")
    parser.add_argument("--files", nargs="+", metavar="FILE",
                        help="Explicit ordered file list (names relative to drafts/). "
                             "Overrides auto-discovery.")
    parser.add_argument("--workspace", default=None,
                        help="Path to federal-proposal-assistant root. "
                             "Defaults to the parent of the tools/ directory.")
    args = parser.parse_args()

    # Resolve paths
    workspace = Path(args.workspace) if args.workspace else WORKSPACE_ROOT
    proposal_dir = workspace / "proposals" / args.proposal
    drafts_dir = proposal_dir / "drafts"
    out_dir = proposal_dir / "final" / "docx"

    if not proposal_dir.exists():
        print(f"ERROR: Proposal directory not found: {proposal_dir}", file=sys.stderr)
        sys.exit(1)
    if not drafts_dir.exists():
        print(f"ERROR: drafts/ directory not found at {drafts_dir}", file=sys.stderr)
        sys.exit(1)

    out_dir.mkdir(parents=True, exist_ok=True)

    # Determine which files to convert
    if args.files:
        md_files = [drafts_dir / f for f in args.files]
        missing = [f for f in md_files if not f.exists()]
        if missing:
            for m in missing:
                print(f"ERROR: File not found: {m}", file=sys.stderr)
            sys.exit(1)
    else:
        md_files = discover_drafts(drafts_dir)
        if not md_files:
            print(f"ERROR: No .md files found in {drafts_dir}", file=sys.stderr)
            sys.exit(1)

    print(f"Proposal : {args.proposal}")
    print(f"Drafts   : {drafts_dir}")
    print(f"Output   : {out_dir}")
    print(f"Mode     : {args.mode}")
    print(f"Files    : {[f.name for f in md_files]}")
    print()

    outputs = []

    # ── Individual files ──────────────────────────────────────────────────────
    if args.mode in ("individual", "both"):
        for md_path in md_files:
            doc = Document()
            setup_document(doc)
            n = convert_md_to_doc(md_path, doc)
            out_path = out_dir / (md_path.stem + ".docx")
            doc.save(out_path)
            outputs.append(out_path)
            print(f"  [OK] {out_path.name}  ({n} lines)")

    # ── Combined document ─────────────────────────────────────────────────────
    if args.mode in ("combined", "both"):
        doc2 = Document()
        setup_document(doc2)
        total = 0
        for idx, md_path in enumerate(md_files):
            n = convert_md_to_doc(md_path, doc2, page_break=(idx > 0))
            total += n
        out_path2 = out_dir / f"{args.combined_name}.docx"
        doc2.save(out_path2)
        outputs.append(out_path2)
        print(f"  [OK] {out_path2.name}  ({total} lines total, {len(md_files)} sections)")

    print()
    print(f"Done — {len(outputs)} file(s) written to {out_dir}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
